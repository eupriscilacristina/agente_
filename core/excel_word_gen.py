# AG-MTE - Excel and Word Report Generator
# Este script cuida da automacao de exportar as analises do agente
# diretamente para arquivos formatados profissionalmente para a diretoria.

import pandas as pd
from docx import Document
from datetime import datetime


class ReportGenerator:
    """
    Classe para gerar relatorios executivos em formato Excel e Word.
    """

    @staticmethod
    def generate_excel(data: list, filename: str = "relatorio_demandas.xlsx") -> str:
        """
        Gera uma planilha formatada pronta para Excel a partir das demandas.

        Args:
            data: Lista de dicionarios com os dados das demandas.
            filename: Nome do arquivo de saida.

        Returns:
            str: Mensagem de confirmacao.
        """
        if not data:
            return "Nenhuma demanda para gerar planilha."

        df = pd.DataFrame(data)

        # Colunas desejadas na ordem
        colunas_desejadas = [
            "id", "titulo", "descricao", "complexidade",
            "prazo_sugerido_dias", "fase", "risco_principal",
            "plano_contingencia", "data_criacao", "status"
        ]

        # Filtrar apenas colunas que existem
        colunas_existentes = [col for col in colunas_desejadas if col in df.columns]
        df = df[colunas_existentes]

        # Renomear colunas para portugues
        mapeamento_colunas = {
            "id": "ID",
            "titulo": "Demanda",
            "descricao": "Descricao",
            "complexidade": "Complexidade",
            "prazo_sugerido_dias": "Prazo (Dias)",
            "fase": "Fase",
            "risco_principal": "Risco Principal",
            "plano_contingencia": "Plano Contingencia",
            "data_criacao": "Data Criacao",
            "status": "Status"
        }
        df = df.rename(columns=mapeamento_colunas)

        # Gerar planilha
        df.to_excel(filename, index=False, engine="openpyxl")
        return f"Planilha '{filename}' gerada com sucesso! ({len(df)} demandas)"

    @staticmethod
    def generate_word(
        exec_summary: str,
        details: str,
        filename: str = "status_report_diretoria.docx"
    ) -> str:
        """
        Gera um relatorio executivo formatado em Word (.docx).

        Args:
            exec_summary: Texto do resumo executivo.
            details: Texto com detalhamento de entregas e riscos.
            filename: Nome do arquivo de saida.

        Returns:
            str: Mensagem de confirmacao.
        """
        doc = Document()

        # Cabecalho Executivo
        doc.add_heading("Relatorio de Status e Gestao de Processos", 0)
        doc.add_paragraph(f"Data de Emissao: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("Departamento de Tecnologia e Processos - NIT\n")

        # Conteudo
        doc.add_heading("1. Resumo Executivo", level=1)
        doc.add_paragraph(exec_summary)

        doc.add_heading("2. Detalhamento de Entregas e Riscos", level=1)
        doc.add_paragraph(details)

        # Rodape
        doc.add_paragraph("\n" + "=" * 50)
        doc.add_paragraph("Documento gerado automaticamente pelo AG-MTE")
        doc.add_paragraph(f"Versao: 1.0 | Emissao: {datetime.now().strftime('%d/%m/%Y')}")

        doc.save(filename)
        return f"Documento Word '{filename}' gerado com sucesso!"
