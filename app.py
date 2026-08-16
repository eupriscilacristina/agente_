import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime
from docx import Document
import plotly.express as px
import plotly.graph_objects as go

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(
    page_title="AG-MTE | Agente Master de Tecnologia e Processos",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- ESTILIZAÇÃO VISUAL ESPETACULAR ---
st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">

<style>
    /* ===== RESET & BASE ===== */
    .stApp {
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        background: linear-gradient(135deg, #0F172A 0%, #1E293B 50%, #0F172A 100%);
    }
    
    /* ===== HEADER PRINCIPAL ===== */
    .hero-container {
        background: linear-gradient(135deg, rgba(37, 99, 235, 0.15) 0%, rgba(14, 165, 233, 0.1) 100%);
        border: 1px solid rgba(37, 99, 235, 0.3);
        border-radius: 24px;
        padding: 40px;
        margin-bottom: 30px;
        position: relative;
        overflow: hidden;
    }
    
    .hero-container::before {
        content: '';
        position: absolute;
        top: -50%;
        right: -50%;
        width: 100%;
        height: 100%;
        background: radial-gradient(circle, rgba(14, 165, 233, 0.1) 0%, transparent 70%);
        animation: pulse 4s ease-in-out infinite;
    }
    
    @keyframes pulse {
        0%, 100% { transform: scale(1); opacity: 0.5; }
        50% { transform: scale(1.1); opacity: 0.8; }
    }
    
    @keyframes fadeInUp {
        from { opacity: 0; transform: translateY(30px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    @keyframes slideInLeft {
        from { opacity: 0; transform: translateX(-30px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    @keyframes glow {
        0%, 100% { box-shadow: 0 0 20px rgba(37, 99, 235, 0.3); }
        50% { box-shadow: 0 0 40px rgba(37, 99, 235, 0.6); }
    }
    
    .hero-title {
        font-family: 'Space Grotesk', sans-serif;
        font-size: 3.2rem;
        font-weight: 700;
        background: linear-gradient(135deg, #FFFFFF 0%, #94A3B8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 12px;
        animation: fadeInUp 0.8s ease-out;
        position: relative;
        z-index: 1;
    }
    
    .hero-subtitle {
        font-size: 1.2rem;
        color: #94A3B8;
        font-weight: 400;
        animation: fadeInUp 0.8s ease-out 0.2s both;
        position: relative;
        z-index: 1;
    }
    
    .hero-badge {
        display: inline-flex;
        align-items: center;
        gap: 8px;
        background: rgba(16, 185, 129, 0.2);
        border: 1px solid rgba(16, 185, 129, 0.4);
        color: #10B981;
        padding: 8px 16px;
        border-radius: 50px;
        font-size: 0.85rem;
        font-weight: 600;
        margin-top: 16px;
        animation: fadeInUp 0.8s ease-out 0.4s both;
        position: relative;
        z-index: 1;
    }
    
    .hero-badge::before {
        content: '';
        width: 8px;
        height: 8px;
        background: #10B981;
        border-radius: 50%;
        animation: blink 1.5s ease-in-out infinite;
    }
    
    @keyframes blink {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.3; }
    }
    
    /* ===== CARDS DE MÉTRICAS ===== */
    .metric-card {
        background: linear-gradient(145deg, rgba(30, 41, 59, 0.8) 0%, rgba(15, 23, 42, 0.9) 100%);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 20px;
        padding: 28px;
        position: relative;
        overflow: hidden;
        transition: all 0.4s cubic-bezier(0.175, 0.885, 0.32, 1.275);
        animation: fadeInUp 0.6s ease-out both;
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, var(--accent-color, #2563EB) 0%, transparent 100%);
    }
    
    .metric-card:hover {
        transform: translateY(-8px) scale(1.02);
        border-color: rgba(255, 255, 255, 0.15);
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.3);
    }
    
    .metric-icon-wrapper {
        width: 56px;
        height: 56px;
        border-radius: 16px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.8rem;
        margin-bottom: 16px;
        background: linear-gradient(135deg, var(--icon-bg, rgba(37, 99, 235, 0.2)) 0%, rgba(37, 99, 235, 0.1) 100%);
    }
    
    .metric-value {
        font-family: 'Space Grotesk', sans-serif;
        font-size: 2.8rem;
        font-weight: 700;
        color: #FFFFFF;
        margin-bottom: 4px;
        line-height: 1;
    }
    
    .metric-label {
        font-size: 0.85rem;
        color: #64748B;
        font-weight: 500;
        text-transform: uppercase;
        letter-spacing: 0.08em;
    }
    
    .metric-change {
        display: inline-flex;
        align-items: center;
        gap: 4px;
        margin-top: 12px;
        padding: 4px 10px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 600;
    }
    
    .metric-change.positive {
        background: rgba(16, 185, 129, 0.2);
        color: #10B981;
    }
    
    .metric-change.neutral {
        background: rgba(148, 163, 184, 0.2);
        color: #94A3B8;
    }
    
    /* ===== CARDS PRINCIPAIS ===== */
    .stCard {
        background: linear-gradient(145deg, rgba(30, 41, 59, 0.6) 0%, rgba(15, 23, 42, 0.8) 100%) !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
        border-radius: 20px !important;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2) !important;
    }
    
    /* ===== SIDEBAR ===== */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #0F172A 0%, #1E293B 100%) !important;
        border-right: 1px solid rgba(255, 255, 255, 0.08) !important;
    }
    
    section[data-testid="stSidebar"] .stRadio label {
        color: #E2E8F0 !important;
        font-weight: 500;
        padding: 12px 16px;
        border-radius: 12px;
        transition: all 0.3s ease;
    }
    
    section[data-testid="stSidebar"] .stRadio label:hover {
        background: rgba(255, 255, 255, 0.05);
    }
    
    section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p {
        color: #94A3B8;
    }
    
    section[data-testid="stSidebar"]::before {
        content: '⚡ AG-MTE';
        display: block;
        padding: 20px 24px;
        font-family: 'Space Grotesk', sans-serif;
        font-size: 1.5rem;
        font-weight: 700;
        color: #FFFFFF;
        border-bottom: 1px solid rgba(255, 255, 255, 0.08);
        margin-bottom: 20px;
    }
    
    /* ===== BOTÕES ===== */
    .stButton > button {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 14px 32px !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275) !important;
        box-shadow: 0 4px 15px rgba(37, 99, 235, 0.4) !important;
        position: relative;
        overflow: hidden;
    }
    
    .stButton > button::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255, 255, 255, 0.2), transparent);
        transition: left 0.5s ease;
    }
    
    .stButton > button:hover::before {
        left: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 8px 25px rgba(37, 99, 235, 0.5) !important;
    }
    
    /* ===== FORMULÁRIOS ===== */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div,
    .stSlider > div > div > div {
        background: rgba(30, 41, 59, 0.6) !important;
        border: 1px solid rgba(255, 255, 255, 0.1) !important;
        border-radius: 12px !important;
        color: #FFFFFF !important;
        padding: 14px 18px !important;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #2563EB !important;
        box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.2) !important;
    }
    
    .stTextInput > div > div > input::placeholder,
    .stTextArea > div > div > textarea::placeholder {
        color: #64748B !important;
    }
    
    /* ===== SLIDER ===== */
    .stSlider > div > div > div > div {
        background: linear-gradient(90deg, #2563EB 0%, #0EA5E9 100%) !important;
    }
    
    /* ===== ALERTAS ===== */
    .stSuccess {
        background: linear-gradient(135deg, rgba(16, 185, 129, 0.15) 0%, rgba(16, 185, 129, 0.05) 100%) !important;
        border: 1px solid rgba(16, 185, 129, 0.3) !important;
        border-radius: 16px !important;
        color: #10B981 !important;
    }
    
    .stWarning {
        background: linear-gradient(135deg, rgba(245, 158, 11, 0.15) 0%, rgba(245, 158, 11, 0.05) 100%) !important;
        border: 1px solid rgba(245, 158, 11, 0.3) !important;
        border-radius: 16px !important;
        color: #F59E0B !important;
    }
    
    .stError {
        background: linear-gradient(135deg, rgba(239, 68, 68, 0.15) 0%, rgba(239, 68, 68, 0.05) 100%) !important;
        border: 1px solid rgba(239, 68, 68, 0.3) !important;
        border-radius: 16px !important;
        color: #EF4444 !important;
    }
    
    .stInfo {
        background: linear-gradient(135deg, rgba(14, 165, 233, 0.15) 0%, rgba(14, 165, 233, 0.05) 100%) !important;
        border: 1px solid rgba(14, 165, 233, 0.3) !important;
        border-radius: 16px !important;
        color: #0EA5E9 !important;
    }
    
    /* ===== TABELAS ===== */
    .stDataFrame {
        border-radius: 16px !important;
        overflow: hidden !important;
        border: 1px solid rgba(255, 255, 255, 0.08) !important;
    }
    
    /* ===== DIVIDERS ===== */
    hr {
        border: none !important;
        border-top: 1px solid rgba(255, 255, 255, 0.08) !important;
        margin: 24px 0 !important;
    }
    
    /* ===== TEXTOS ===== */
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Space Grotesk', sans-serif !important;
        color: #FFFFFF !important;
    }
    
    p, li, span, label {
        color: #CBD5E1 !important;
    }
    
    /* ===== EXPANDER ===== */
    .streamlit-expanderHeader {
        background: rgba(30, 41, 59, 0.6) !important;
        border-radius: 12px !important;
        font-weight: 600 !important;
    }
    
    /* ===== TABS ===== */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: rgba(30, 41, 59, 0.6);
        padding: 8px;
        border-radius: 16px;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 12px !important;
        padding: 12px 24px !important;
        font-weight: 500 !important;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        color: white !important;
    }
    
    /* ===== RESULTADO CARD ===== */
    .result-card {
        background: linear-gradient(145deg, rgba(14, 165, 233, 0.1) 0%, rgba(37, 99, 235, 0.05) 100%);
        border: 1px solid rgba(14, 165, 233, 0.3);
        border-left: 5px solid #0EA5E9;
        padding: 32px;
        border-radius: 16px;
        margin-top: 24px;
        animation: fadeInUp 0.6s ease-out;
    }
    
    .result-card h3 {
        color: #FFFFFF;
        font-size: 1.4rem;
        font-weight: 700;
        margin-bottom: 24px;
        display: flex;
        align-items: center;
        gap: 12px;
    }
    
    .result-item {
        display: flex;
        align-items: flex-start;
        margin-bottom: 16px;
        padding: 20px;
        background: rgba(255, 255, 255, 0.03);
        border-radius: 12px;
        border: 1px solid rgba(255, 255, 255, 0.05);
        transition: all 0.3s ease;
    }
    
    .result-item:hover {
        background: rgba(255, 255, 255, 0.06);
        border-color: rgba(255, 255, 255, 0.1);
        transform: translateX(5px);
    }
    
    .result-icon {
        font-size: 1.8rem;
        margin-right: 16px;
        flex-shrink: 0;
    }
    
    .result-content {
        flex: 1;
    }
    
    .result-label {
        font-weight: 600;
        color: #94A3B8;
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        margin-bottom: 6px;
    }
    
    .result-value {
        color: #FFFFFF;
        font-size: 1.05rem;
        line-height: 1.6;
    }
    
    /* ===== PROGRESS BAR ===== */
    .progress-bar {
        height: 8px;
        background: rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        overflow: hidden;
        margin-top: 12px;
    }
    
    .progress-fill {
        height: 100%;
        background: linear-gradient(90deg, #2563EB 0%, #0EA5E9 100%);
        border-radius: 10px;
        transition: width 0.5s ease;
    }
    
    /* ===== FOOTER ===== */
    .footer {
        text-align: center;
        padding: 30px;
        margin-top: 40px;
        border-top: 1px solid rgba(255, 255, 255, 0.08);
        color: #64748B;
        font-size: 0.85rem;
    }
    
    .footer-brand {
        font-family: 'Space Grotesk', sans-serif;
        font-weight: 700;
        color: #94A3B8;
    }
</style>
""", unsafe_allow_html=True)

# --- BANCO DE DADOS LOCAL (JSON) ---
DB_PATH = "data/demands_db.json"

def carregar_dados():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    if not os.path.exists(DB_PATH):
        with open(DB_PATH, "w", encoding="utf-8") as f:
            json.dump([], f)
    with open(DB_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def salvar_dados(dados):
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    with open(DB_PATH, "w", encoding="utf-8") as f:
        json.dump(dados, f, ensure_ascii=False, indent=4)

# --- MOTOR DE GERAÇÃO DE RELATÓRIOS ---
def gerar_word(demanda):
    doc = Document()
    doc.add_heading('Relatório Executivo de Status - AG-MTE', 0)
    doc.add_paragraph(f"Emitido em: {datetime.now().strftime('%d/%m/%Y %H:%M')} | Departamento de Tecnologia e Processos")
    
    doc.add_heading('1. Visão Geral da Demanda', level=1)
    doc.add_paragraph(f"Título: {demanda['titulo']}")
    doc.add_paragraph(f"Descrição: {demanda['descricao']}")
    doc.add_paragraph(f"Fase Atual: {demanda['fase']}")
    
    doc.add_heading('2. Cronograma e SLA Preditivo', level=1)
    doc.add_paragraph(f"Prazo Estimado: {demanda['prazo_sugerido_dias']} dias úteis")
    
    doc.add_heading('3. Gestão de Riscos e Plano de Contingência', level=1)
    doc.add_paragraph(f"Risco Mapeado: {demanda['risco_principal']}")
    doc.add_paragraph(f"Plano de Mitigação: {demanda['plano_contingencia']}")
    
    filename = f"relatorio_{demanda['id']}.docx"
    doc.save(filename)
    return filename

def gerar_excel(demandas):
    df = pd.DataFrame(demandas)
    filename = "relatorio_demandas_completo.xlsx"
    df.to_excel(filename, index=False, engine='openpyxl')
    return filename

# --- DADOS PARA GRÁFICOS ---
def criar_grafico_complexidade(demandas):
    if not demandas:
        return None
    df = pd.DataFrame(demandas)
    fig = px.pie(df, names='complexidade', title='Distribuição por Complexidade',
                 color_discrete_sequence=['#10B981', '#3B82F6', '#F59E0B', '#F97316', '#EF4444'])
    fig.update_layout(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#94A3B8'),
        showlegend=True,
        legend=dict(font=dict(color='#94A3B8'))
    )
    return fig

def criar_grafico_prazos(demandas):
    if not demandas:
        return None
    df = pd.DataFrame(demandas)
    fig = px.bar(df, x='titulo', y='prazo_sugerido_dias', title='Prazos por Demanda',
                 color='complexidade', color_continuous_scale='Blues')
    fig.update_layout(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#94A3B8'),
        xaxis=dict(gridcolor='rgba(255,255,255,0.05)'),
        yaxis=dict(gridcolor='rgba(255,255,255,0.05)')
    )
    return fig

def criar_grafico_timeline(demandas):
    if not demandas:
        return None
    df = pd.DataFrame(demandas)
    fig = px.timeline(df, x_start='data_criacao', x_end='data_criacao', y='titulo', 
                      color='complexidade', title='Timeline de Demandas',
                      color_continuous_scale='Viridis')
    fig.update_layout(
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#94A3B8')
    )
    return fig

# --- INTERFACE PRINCIPAL ---

# Sidebar
with st.sidebar:
    menu = st.radio(
        "Navegação:",
        ["🏠 Dashboard Executivo", "🚀 Nova Demanda & Análise", "📅 Acompanhamento Diário", "📊 Painel de Controle", "📑 Central de Relatórios"],
        label_visibility="collapsed"
    )
    st.markdown("<br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br><br>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("""
    <div style='padding: 16px; background: rgba(255,255,255,0.03); border-radius: 12px; border: 1px solid rgba(255,255,255,0.08);'>
        <p style='margin: 0; line-height: 1.6; color: #64748B; font-size: 0.75rem;'>
            <strong style='color: #FFFFFF; font-size: 0.9rem; font-family: Space Grotesk, sans-serif;'>⚡ AG-MTE v2.0</strong><br><br>
            <span style='color: #475569;'>© 2026 Priscila Cristina.<br>Todos os direitos reservados.</span>
        </p>
    </div>
    """, unsafe_allow_html=True)

demandas = carregar_dados()

# Header com nome completo centralizado e status online
st.markdown("""
<div style='text-align: center; padding: 20px 24px; background: linear-gradient(135deg, rgba(37, 99, 235, 0.1) 0%, rgba(14, 165, 233, 0.05) 100%); border: 1px solid rgba(37, 99, 235, 0.2); border-radius: 12px; margin-bottom: 24px;'>
    <p style='margin: 0 0 8px 0; font-size: 1.3rem; font-weight: 700; color: #FFFFFF; letter-spacing: 0.02em; font-family: Space Grotesk, sans-serif;'>⚡ AG-MTE <span style='color: #64748B; font-weight: 300;'>|</span> <span style='color: #94A3B8; font-weight: 400;'>Agente Master de Tecnologia e Processos</span></p>
    <div style='display: flex; align-items: center; justify-content: center; gap: 8px;'>
        <span style='width: 8px; height: 8px; background: #10B981; border-radius: 50%; box-shadow: 0 0 8px #10B981;'></span>
        <span style='color: #10B981; font-size: 0.85rem; font-weight: 600;'>Sistema Online</span>
    </div>
</div>
""", unsafe_allow_html=True)

# --- MÓDULO 1: DASHBOARD EXECUTIVO ---
if menu == "🏠 Dashboard Executivo":
    st.markdown("## 📊 Dashboard Executivo")
    st.markdown("Visão geral do sistema de gestão de demandas e entregas.")
    
    if demandas:
        df = pd.DataFrame(demandas)
        
        # Métricas Principais
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.markdown("""
            <div class="metric-card" style="--accent-color: #3B82F6; --icon-bg: rgba(59, 130, 246, 0.2);">
                <div class="metric-icon-wrapper">📋</div>
                <div class="metric-value">""" + str(len(demandas)) + """</div>
                <div class="metric-label">Total de Demandas</div>
                <div class="metric-change positive">↑ Ativo</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            media_prazo = df["prazo_sugerido_dias"].mean()
            st.markdown(f"""
            <div class="metric-card" style="--accent-color: #0EA5E9; --icon-bg: rgba(14, 165, 233, 0.2);">
                <div class="metric-icon-wrapper">⏱️</div>
                <div class="metric-value">{media_prazo:.1f}</div>
                <div class="metric-label">Média de Prazo</div>
                <div class="metric-change neutral">dias úteis</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            media_complexidade = df["complexidade"].mean()
            st.markdown(f"""
            <div class="metric-card" style="--accent-color: #F59E0B; --icon-bg: rgba(245, 158, 11, 0.2);">
                <div class="metric-icon-wrapper">🎯</div>
                <div class="metric-value">{media_complexidade:.1f}</div>
                <div class="metric-label">Complexidade Média</div>
                <div class="metric-change neutral">/ 5.0</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col4:
            st.markdown("""
            <div class="metric-card" style="--accent-color: #10B981; --icon-bg: rgba(16, 185, 129, 0.2);">
                <div class="metric-icon-wrapper">✅</div>
                <div class="metric-value">100%</div>
                <div class="metric-label">Status Operacional</div>
                <div class="metric-change positive">● Online</div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Gráficos
        col_graf1, col_graf2 = st.columns(2)
        
        with col_graf1:
            fig_complexidade = criar_grafico_complexidade(demandas)
            if fig_complexidade:
                st.plotly_chart(fig_complexidade, use_container_width=True)
        
        with col_graf2:
            fig_prazos = criar_grafico_prazos(demandas)
            if fig_prazos:
                st.plotly_chart(fig_prazos, use_container_width=True)
        
        # Últimas Demandas
        st.markdown("---")
        st.markdown("### 📋 Últimas Demandas Cadastradas")
        st.dataframe(df.tail(5), use_container_width=True)
        
    else:
        st.info("📭 Nenhuma demanda cadastrada ainda. Comece adicionando uma nova demanda!")

# --- MÓDULO 2: NOVA DEMANDA ---
elif menu == "🚀 Nova Demanda & Análise":
    st.markdown("## 🚀 Cadastro de Nova Demanda")
    st.markdown("Insira os dados para que o AG-MTE calcule prazos, mapeie riscos e estruture o fluxo.")
    
    with st.form("form_demanda"):
        col1, col2 = st.columns([3, 1])
        with col1:
            titulo = st.text_input("📋 Título do Projeto / Demanda:", placeholder="Ex: Automação de Relatórios de TI")
        with col2:
            complexidade = st.selectbox("🎯 Complexidade:", [1, 2, 3, 4, 5], index=2, 
                                       format_func=lambda x: f"{x} - {'Muito Baixa' if x==1 else 'Baixa' if x==2 else 'Média' if x==3 else 'Alta' if x==4 else 'Crítica'}")
            
        descricao = st.text_area("📝 Escopo Detalhado / Requisitos:", 
                                 placeholder="Descreva o que foi pedido, os objetivos e restrições...", 
                                 height=140)
        
        st.markdown("---")
        btn_executar = st.form_submit_button("🤖 Processar Demanda com IA", use_container_width=True)
        
        if btn_executar:
            if not titulo.strip():
                st.error("❌ O título da demanda é obrigatório.")
            else:
                # Lógica preditiva avançada
                dias_estimados = complexidade * 3
                
                riscos = {
                    1: "Risco mínimo. Tarefa operacional padrão sem dependências.",
                    2: "Baixo risco. Possíveis ajustes menores em processos.",
                    3: "Risco moderado. Dependências externas podem impactar prazo.",
                    4: "Risco alto. Possível gargalo técnico ou dependência cruzada de infraestrutura.",
                    5: "Risco crítico. Alta complexidade com múltiplas dependências e impacto sistêmico."
                }
                
                contingencias = {
                    1: "Execução direta com validação pontual.",
                    2: "Revisão incremental a cada 2 dias.",
                    3: "Validação semanal e ajuste ágil de escopo.",
                    4: "Reuniões de alinhamento a cada 3 dias. Plano B para integrações.",
                    5: "Comitê de crise. Revisão diária. Escopo mínimo viável (MVP)."
                }
                
                fases = {
                    1: "Execução Rápida",
                    2: "Planejamento / Início",
                    3: "Desenvolvimento / Execução",
                    4: "Desenvolvimento / Validação",
                    5: "Planejamento Estratégico / Execução Completa"
                }
                
                risco = riscos.get(complexidade, "Risco não mapeado.")
                contingencia = contingencias.get(complexidade, "Plano padrão.")
                fase = fases.get(complexidade, "Planejamento / Início")
                
                nova = {
                    "id": len(demandas) + 1,
                    "titulo": titulo,
                    "descricao": descricao,
                    "complexidade": complexidade,
                    "prazo_sugerido_dias": dias_estimados,
                    "fase": fase,
                    "risco_principal": risco,
                    "plano_contingencia": contingencia,
                    "data_criacao": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
                
                demandas.append(nova)
                salvar_dados(demandas)
                
                st.success("✅ Demanda processada, planejada e salva com sucesso!")
                
                # Barra de progresso animada
                progresso = min(100, (complexidade / 5) * 100)
                st.markdown(f"""
                <div class="progress-bar">
                    <div class="progress-fill" style="width: {progresso}%;"></div>
                </div>
                """, unsafe_allow_html=True)
                
                # Resultado formatado
                st.markdown(f"""
                <div class="result-card">
                    <h3>📋 Resultado da Análise de Engenharia</h3>
                    <div class="result-item">
                        <span class="result-icon">⏱️</span>
                        <div class="result-content">
                            <div class="result-label">Prazo Adequado (SLA)</div>
                            <div class="result-value">{dias_estimados} dias úteis</div>
                        </div>
                    </div>
                    <div class="result-item">
                        <span class="result-icon">📊</span>
                        <div class="result-content">
                            <div class="result-label">Fase do Projeto</div>
                            <div class="result-value">{fase}</div>
                        </div>
                    </div>
                    <div class="result-item">
                        <span class="result-icon">⚠️</span>
                        <div class="result-content">
                            <div class="result-label">Risco Identificado</div>
                            <div class="result-value">{risco}</div>
                        </div>
                    </div>
                    <div class="result-item">
                        <span class="result-icon">🛡️</span>
                        <div class="result-content">
                            <div class="result-label">Plano de Contingência</div>
                            <div class="result-value">{contingencia}</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

# --- MÓDULO 3: PAINEL DE CONTROLE ---
elif menu == "📊 Painel de Controle":
    st.markdown("## 📊 Painel de Controle de Entregas")
    st.markdown("Acompanhe o andamento de tudo o que foi iniciado.")
    
    if not demandas:
        st.info("📭 Nenhuma demanda cadastrada ainda.")
    else:
        df = pd.DataFrame(demandas)
        
        # Métricas
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📋 Total", len(demandas))
        with col2:
            st.metric("⏱️ Média Prazo", f"{df['prazo_sugerido_dias'].mean():.1f} dias")
        with col3:
            st.metric("🎯 Complexidade Média", f"{df['complexidade'].mean():.1f}/5")
        with col4:
            st.metric("✅ Status", "Online")
        
        st.markdown("---")
        
        # Filtros
        st.markdown("### 🔍 Filtros")
        col_f1, col_f2 = st.columns(2)
        with col_f1:
            filtro_complexidade = st.multiselect("Filtrar por Complexidade:", [1, 2, 3, 4, 5], default=[1, 2, 3, 4, 5])
        with col_f2:
            filtro_fase = st.multiselect("Filtrar por Fase:", df['fase'].unique().tolist(), default=df['fase'].unique().tolist())
        
        # Aplicar filtros
        df_filtrado = df[(df['complexidade'].isin(filtro_complexidade)) & (df['fase'].isin(filtro_fase))]
        
        st.markdown(f"### 📋 Demandas ({len(df_filtrado)} registros)")
        st.dataframe(df_filtrado, use_container_width=True, height=400)
        
        # Gráfico de timeline
        if len(df) > 1:
            st.markdown("### 📈 Evolução dos Prazos")
            fig_evolucao = px.line(df, x='data_criacao', y='prazo_sugerido_dias', 
                                  title='Evolução dos Prazos ao Longo do Tempo',
                                  markers=True, color_discrete_sequence=['#3B82F6'])
            fig_evolucao.update_layout(
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='#94A3B8')
            )
            st.plotly_chart(fig_evolucao, use_container_width=True)

# --- MÓDULO 4: RELATÓRIOS ---
elif menu == "📑 Central de Relatórios":
    st.markdown("## 📑 Central de Relatórios Executivos")
    st.markdown("Gere documentos formatados para apresentação à diretoria.")
    
    if not demandas:
        st.warning("⚠️ Cadastre demandas para gerar relatórios.")
    else:
        titulos_demandas = [d['titulo'] for d in demandas]
        escolha = st.selectbox("🔍 Selecione a Demanda:", titulos_demandas)
        
        demanda_selecionada = next(d for d in demandas if d['titulo'] == escolha)
        
        # Card da demanda selecionada
        st.markdown(f"""
        <div class="result-card">
            <h3>📄 Demanda Selecionada</h3>
            <div class="result-item">
                <span class="result-icon">📋</span>
                <div class="result-content">
                    <div class="result-label">Título</div>
                    <div class="result-value">{demanda_selecionada['titulo']}</div>
                </div>
            </div>
            <div class="result-item">
                <span class="result-icon">📊</span>
                <div class="result-content">
                    <div class="result-label">Complexidade</div>
                    <div class="result-value">{demanda_selecionada['complexidade']}/5</div>
                </div>
            </div>
            <div class="result-item">
                <span class="result-icon">⏱️</span>
                <div class="result-content">
                    <div class="result-label">Prazo Estimado</div>
                    <div class="result-value">{demanda_selecionada['prazo_sugerido_dias']} dias úteis</div>
                </div>
            </div>
            <div class="result-item">
                <span class="result-icon">📅</span>
                <div class="result-content">
                    <div class="result-label">Data de Criação</div>
                    <div class="result-value">{demanda_selecionada['data_criacao']}</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        col_w, col_e = st.columns(2)
        
        with col_w:
            st.markdown("### 📄 Documento Executivo (Word)")
            st.markdown("Gere um relatório formatado em `.docx` pronto para apresentação.")
            if st.button("📥 Gerar e Baixar Relatório Word", use_container_width=True):
                arquivo_word = gerar_word(demanda_selecionada)
                with open(arquivo_word, "rb") as f:
                    st.download_button(
                        label="📥 Clique para Baixar Word",
                        data=f,
                        file_name=arquivo_word,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                st.success("✅ Relatório compilado com sucesso!")
                    
        with col_e:
            st.markdown("### 📊 Planilha de Controle (Excel)")
            st.markdown("Exporte todos os dados em formato `.xlsx` para análise.")
            if st.button("📥 Gerar e Baixar Planilha Excel", use_container_width=True):
                arquivo_excel = gerar_excel(demandas)
                with open(arquivo_excel, "rb") as f:
                    st.download_button(
                        label="📥 Clique para Baixar Excel",
                        data=f,
                        file_name=arquivo_excel,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                    st.success("✅ Planilha gerada com sucesso!")

# --- MÓDULO 5: ACOMPANHAMENTO DIÁRIO ---
elif menu == "📅 Acompanhamento Diário":
    st.markdown("## 📅 Acompanhamento Diário e Ciclo de Vida")
    st.markdown("Registre o progresso diário de cada demanda, documente avanços e acompanhe o histórico completo.")

    LOG_PATH = "data/daily_logs.json"

    def carregar_logs():
        os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
        if not os.path.exists(LOG_PATH):
            with open(LOG_PATH, "w", encoding="utf-8") as f:
                json.dump({}, f)
        with open(LOG_PATH, "r", encoding="utf-8") as f:
            return json.load(f)

    def salvar_logs(logs):
        os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
        with open(LOG_PATH, "w", encoding="utf-8") as f:
            json.dump(logs, f, ensure_ascii=False, indent=4)

    logs = carregar_logs()

    if not demandas:
        st.info("📭 Cadastre uma demanda primeiro no menu '🚀 Nova Demanda & Análise'.")
    else:
        titulos = [d["titulo"] for d in demandas]
        ids = [d["id"] for d in demandas]
        mapa_id = {d["titulo"]: d["id"] for d in demandas}

        col_sel1, col_sel2 = st.columns([3, 1])
        with col_sel1:
            titulo_sel = st.selectbox("🔍 Selecione a Demanda / Projeto:", titulos)
        with col_sel2:
            id_sel = mapa_id[titulo_sel]
            demanda_atual = next(d for d in demandas if d["id"] == id_sel)
            status_proj = demanda_atual.get("status_projeto", "Em Andamento")
            cor_status = "#10B981" if status_proj == "Finalizado" else "#F59E0B"
            st.markdown(f"""
            <div style='text-align:center; padding:14px; background:rgba(255,255,255,0.03); border-radius:12px; border:1px solid rgba(255,255,255,0.08); margin-top:28px;'>
                <span style='font-size:0.75rem; color:#64748B; text-transform:uppercase; letter-spacing:0.08em;'>Status</span><br>
                <span style='font-size:1.1rem; font-weight:700; color:{cor_status};'>{status_proj}</span>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("---")

        # --- FORMULÁRIO DE REGISTRO DIÁRIO ---
        st.markdown("### ✏️ Registrar Andamento do Dia")

        with st.form("form_daily_log"):
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                data_log = st.date_input("📅 Data:", value=datetime.now().date())
            with col_d2:
                status_dia = st.selectbox("📌 Status do Dia:", [
                    "Em Andamento",
                    "Bloqueado / Gargalo",
                    "Concluído Parcialmente",
                    "Finalizado"
                ])

            nota_progresso = st.text_area("📝 Nota de Progresso:", placeholder="Descreva o que foi feito, entregue ou ajustado hoje...", height=140)

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                finalizar_projeto = st.checkbox("✅ Marcar projeto como Finalizado com Sucesso")
            with col_btn2:
                st.markdown("")  # espaçador

            btn_salvar = st.form_submit_button("💾 Salvar Registro Diário", use_container_width=True)

            if btn_salvar:
                if not nota_progresso.strip():
                    st.error("❌ Preencha a nota de progresso.")
                else:
                    str_id = str(id_sel)
                    if str_id not in logs:
                        logs[str_id] = []

                    registro = {
                        "data": data_log.strftime("%Y-%m-%d"),
                        "status_dia": status_dia,
                        "nota": nota_progresso.strip(),
                        "finalizado": finalizar_projeto,
                        "registro_em": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    logs[str_id].append(registro)

                    if finalizar_projeto:
                        demanda_atual["status_projeto"] = "Finalizado"
                        demanda_atual["data_finalizacao"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        salvar_dados(demandas)

                    salvar_logs(logs)
                    st.success("✅ Registro diário salvo com sucesso!")
                    if finalizar_projeto:
                        st.success("🎉 Projeto marcado como FINALIZADO!")

        st.markdown("---")

        # --- HISTÓRICO / TIMELINE ---
        st.markdown(f"### 📜 Histórico de Andamento — {titulo_sel}")

        str_id = str(id_sel)
        if str_id in logs and logs[str_id]:
            registros = logs[str_id]

            total_registros = len(registros)
            registros_finalizados = sum(1 for r in registros if r.get("finalizado"))
            total_notas = len(registros)

            col_h1, col_h2, col_h3 = st.columns(3)
            with col_h1:
                st.markdown(f"""
                <div class="metric-card" style="--accent-color: #3B82F6; --icon-bg: rgba(59, 130, 246, 0.2);">
                    <div class="metric-icon-wrapper">📝</div>
                    <div class="metric-value">{total_registros}</div>
                    <div class="metric-label">Registros Totais</div>
                </div>
                """, unsafe_allow_html=True)
            with col_h2:
                st.markdown(f"""
                <div class="metric-card" style="--accent-color: #10B981; --icon-bg: rgba(16, 185, 129, 0.2);">
                    <div class="metric-icon-wrapper">✅</div>
                    <div class="metric-value">{registros_finalizados}</div>
                    <div class="metric-label">Finalizações</div>
                </div>
                """, unsafe_allow_html=True)
            with col_h3:
                dias_desde_inicio = (datetime.now().date() - datetime.strptime(registros[0]["data"], "%Y-%m-%d").date()).days
                st.markdown(f"""
                <div class="metric-card" style="--accent-color: #F59E0B; --icon-bg: rgba(245, 158, 11, 0.2);">
                    <div class="metric-icon-wrapper">⏱️</div>
                    <div class="metric-value">{dias_desde_inicio}</div>
                    <div class="metric-label">Dias Desde Início</div>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("---")

            for i, reg in enumerate(reversed(registros)):
                status_cor_map = {
                    "Em Andamento": "#3B82F6",
                    "Bloqueado / Gargalo": "#EF4444",
                    "Concluído Parcialmente": "#F59E0B",
                    "Finalizado": "#10B981"
                }
                cor_stat = status_cor_map.get(reg["status_dia"], "#94A3B8")
                icone_final = "🎉" if reg.get("finalizado") else "📌"
                label_final = " — FINALIZADO" if reg.get("finalizado") else ""

                st.markdown(f"""
                <div style='padding:20px; background:rgba(255,255,255,0.03); border-radius:14px; border:1px solid rgba(255,255,255,0.08); border-left:4px solid {cor_stat}; margin-bottom:14px;'>
                    <div style='display:flex; justify-content:space-between; align-items:center; margin-bottom:10px;'>
                        <span style='font-weight:700; color:#FFFFFF; font-size:0.95rem;'>{icone_final} Dia {reg['data']}{label_final}</span>
                        <span style='padding:4px 12px; border-radius:20px; font-size:0.75rem; font-weight:600; background:rgba(255,255,255,0.06); color:{cor_stat}; border:1px solid {cor_stat}30;'>{reg['status_dia']}</span>
                    </div>
                    <p style='margin:0; color:#CBD5E1; font-size:0.9rem; line-height:1.6;'>{reg['nota']}</p>
                    <p style='margin:8px 0 0 0; color:#475569; font-size:0.75rem;'>Registrado em: {reg['registro_em']}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("📭 Nenhum registro diário para esta demanda ainda. Use o formulário acima para registrar o primeiro andamento.")
